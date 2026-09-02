import streamlit as st
from datetime import date
from io import BytesIO
from PIL import Image, ImageOps
from xml.sax.saxutils import escape
from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


# =========================================================
# PAGE CONFIGURATION
# =========================================================

st.set_page_config(
    page_title="Quality Rejection Report",
    page_icon="📋",
    layout="wide"
)


# =========================================================
# CUSTOM CSS
# =========================================================

st.markdown("""
<style>

    /* Main page */
    .main {
        padding-top: 20px;
    }

    /* Title */
    .report-title {
        text-align: center;
        font-size: 28px;
        font-weight: 700;
        margin-bottom: 25px;
    }

    /* Row container */
    .row-container {
        border: 1px solid #555;
        margin-bottom: -1px;
        min-height: 70px;
    }

    /* Header cell */
    .header-cell {
        font-weight: 700;
        font-size: 15px;
        padding: 18px 12px;
        min-height: 70px;
        display: flex;
        align-items: center;
    }

    /* Normal input cell */
    .input-cell {
        padding: 8px 10px;
    }

    /* Section heading */
    .section-heading {
        font-weight: 700;
        font-size: 18px;
        margin-top: 25px;
        margin-bottom: 10px;
    }

    /* Red rejection text */
    .rejection-title {
        color: #9b1c1c;
        font-weight: 700;
        font-size: 16px;
        margin-bottom: 10px;
    }

    /* Submit button */
    div.stButton > button {
        width: 100%;
        height: 45px;
        font-weight: 700;
        font-size: 16px;
    }

</style>
""", unsafe_allow_html=True)


# =========================================================
# TITLE
# =========================================================

st.markdown(
    '<div class="report-title">QUALITY REJECTION REPORT</div>',
    unsafe_allow_html=True
)


# =========================================================
# HELPER FUNCTION
# =========================================================

def create_row(label, widget_type="text", key=None, height=None):

    """
    Creates a 2-column row:
    Column 1 = Header
    Column 2 = Input
    """

    col1, col2 = st.columns([1, 1.45])

    with col1:
        st.markdown(
            f'<div class="header-cell">{label}</div>',
            unsafe_allow_html=True
        )

    with col2:

        if widget_type == "text":
            value = st.text_input(
                label,
                key=key,
                label_visibility="collapsed"
            )

        elif widget_type == "textarea":
            value = st.text_area(
                label,
                key=key,
                height=height if height else 120,
                label_visibility="collapsed"
            )

        elif widget_type == "date":
            value = st.date_input(
                label,
                value=date.today(),
                key=key,
                label_visibility="collapsed"
            )

        return value


def create_word_report(report_data, photos):
    document = Document()
    document.add_heading("QUALITY REJECTION REPORT", 0)

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in report_data.items():
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        if label == "Reason of Rejection":
            for photo in photos:
                image = Image.open(BytesIO(photo["data"]))
                image = ImageOps.exif_transpose(image)
                if image.mode not in ("RGB", "RGBA"):
                    image = image.convert("RGBA")
                image_data = BytesIO()
                image.save(image_data, format="PNG")
                image_data.seek(0)
                cells[1].add_paragraph().add_run().add_picture(
                    image_data, width=Inches(3.0)
                )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def create_pdf_report(report_data, photos):
    output = BytesIO()
    document = SimpleDocTemplate(output, pagesize=A4,
                                 rightMargin=0.55 * inch,
                                 leftMargin=0.55 * inch,
                                 topMargin=0.55 * inch,
                                 bottomMargin=0.55 * inch)
    styles = getSampleStyleSheet()
    content = [Paragraph("QUALITY REJECTION REPORT", styles["Title"]), Spacer(0.1, 0.2 * inch)]

    table_data = []
    for label, value in report_data.items():
        right_cell = [Paragraph(escape(str(value)), styles["BodyText"])]
        if label == "Reason of Rejection":
            for photo in photos:
                right_cell.extend([
                    PdfImage(BytesIO(photo["data"]), width=4.3 * inch,
                             height=3.0 * inch, kind="proportional")
                ])
        table_data.append([Paragraph(escape(str(label)), styles["BodyText"]), right_cell])
    table = Table(table_data, colWidths=[2.1 * inch, 4.8 * inch])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8eef5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    content.extend([table, Spacer(1, 0.25 * inch)])

    document.build(content)
    return output.getvalue()


# =========================================================
# BASIC DETAILS
# =========================================================

st.markdown(
    '<div class="section-heading">Report Details</div>',
    unsafe_allow_html=True
)


# ---------------------------------------------------------
# ITEM DESCRIPTION
# ---------------------------------------------------------

col1, col2 = st.columns([1, 1.45])

with col1:
    st.markdown(
        '<div class="header-cell">Item Description</div>',
        unsafe_allow_html=True
    )

with col2:
    item_description = st.text_area(
        "Item Description",
        placeholder="Enter item description...",
        height=100,
        label_visibility="collapsed"
    )


# ---------------------------------------------------------
# ITEM CODE
# ---------------------------------------------------------

col1, col2 = st.columns([1, 1.45])

with col1:
    st.markdown(
        '<div class="header-cell">Item Code</div>',
        unsafe_allow_html=True
    )

with col2:
    item_code = st.text_area(
        "Item Code",
        placeholder="Enter item code(s)...",
        height=80,
        label_visibility="collapsed"
    )


# ---------------------------------------------------------
# LOT QTY
# ---------------------------------------------------------

col1, col2 = st.columns([1, 1.45])

with col1:
    st.markdown(
        '<div class="header-cell">Lot Qty.</div>',
        unsafe_allow_html=True
    )

with col2:
    lot_qty = st.text_area(
        "Lot Quantity",
        placeholder="Example:\n64000 NOS\n59860 NOS",
        height=80,
        label_visibility="collapsed"
    )


# ---------------------------------------------------------
# CHALLAN
# ---------------------------------------------------------

col1, col2 = st.columns([1, 1.45])

with col1:
    st.markdown(
        '<div class="header-cell">Challan No. & Date</div>',
        unsafe_allow_html=True
    )

with col2:

    challan_col1, challan_col2 = st.columns(2)

    with challan_col1:
        challan_no = st.text_input(
            "Challan Number",
            placeholder="Challan No.",
            label_visibility="collapsed"
        )

    with challan_col2:
        challan_date = st.date_input(
            "Challan Date",
            value=date.today(),
            label_visibility="collapsed"
        )


# ---------------------------------------------------------
# GRIN
# ---------------------------------------------------------

col1, col2 = st.columns([1, 1.45])

with col1:
    st.markdown(
        '<div class="header-cell">GRIN No. & Date</div>',
        unsafe_allow_html=True
    )

with col2:

    grin_col1, grin_col2 = st.columns(2)

    with grin_col1:
        grin_no = st.text_input(
            "GRIN Number",
            placeholder="GRIN No.",
            label_visibility="collapsed"
        )

    with grin_col2:
        grin_date = st.date_input(
            "GRIN Date",
            value=date.today(),
            label_visibility="collapsed"
        )


# =========================================================
# REASON FOR REJECTION
# =========================================================

st.markdown(
    '<div class="section-heading">Rejection Details</div>',
    unsafe_allow_html=True
)

col1, col2 = st.columns([1, 1.45])

with col1:

    st.markdown(
        '<div class="header-cell" style="min-height:500px;">'
        'Reason of Rejection With Photo.'
        '</div>',
        unsafe_allow_html=True
    )

with col2:

    st.markdown(
        '<div class="rejection-title">'
        'Reason for Rejection'
        '</div>',
        unsafe_allow_html=True
    )

    rejection_reason = st.text_area(
        "Reason for Rejection",
        placeholder=(
            "Example:\n"
            "Flashes on washer & gate point not cut properly"
        ),
        height=120,
        label_visibility="collapsed"
    )

    st.markdown("### Upload Photos")

    st.caption(
        "Upload one or more photographs showing the rejection/defect."
    )

    uploaded_photos = st.file_uploader(
        "Upload rejection photographs",
        type=["jpg", "jpeg", "png", "webp"],
        accept_multiple_files=True,
        key="rejection_photos"
    )

    # -----------------------------------------------------
    # PHOTO PREVIEW
    # -----------------------------------------------------

    if uploaded_photos:

        st.markdown("#### Photo Preview")

        # Keep every photo inside the right-hand report cell.
        for photo in uploaded_photos:
            image = Image.open(photo)
            st.image(
                image,
                use_column_width=True
            )


# =========================================================
# CORRECTIVE ACTION
# =========================================================

col1, col2 = st.columns([1, 1.45])

with col1:

    st.markdown(
        '<div class="header-cell" style="min-height:180px;">'
        'Corrective Action Taken By External Provider '
        '& Date Of Implementation'
        '</div>',
        unsafe_allow_html=True
    )

with col2:

    corrective_action = st.text_area(
        "Corrective Action",
        placeholder="Enter corrective action taken...",
        height=100,
        label_visibility="collapsed"
    )

    implementation_date = st.date_input(
        "Date of Implementation",
        value=date.today()
    )


# =========================================================
# SUBMIT
# =========================================================

st.markdown("---")

submit = st.button(
    "SUBMIT REJECTION REPORT",
    type="primary"
)


# =========================================================
# SUBMISSION
# =========================================================

if submit:

    if not item_description:
        st.error("Please enter Item Description.")

    elif not item_code:
        st.error("Please enter Item Code.")

    elif not rejection_reason:
        st.error("Please enter Reason for Rejection.")

    else:

        # -------------------------------------------------
        # CREATE DATA OBJECT
        # -------------------------------------------------

        report_data = {
            "Item Description": item_description,
            "Item Code": item_code,
            "Lot Qty": lot_qty,
            "Challan No": challan_no,
            "Challan Date": str(challan_date),
            "GRIN No": grin_no,
            "GRIN Date": str(grin_date),
            "Reason of Rejection": rejection_reason,
            "Corrective Action": corrective_action,
            "Implementation Date": str(implementation_date),
            "Number of Photos": len(uploaded_photos)
        }

        report_photos = [
            {"name": photo.name, "data": photo.getvalue()}
            for photo in uploaded_photos
        ]

        st.session_state["report_data"] = report_data
        st.session_state["report_photos"] = report_photos

        st.success("Rejection report submitted successfully.")


# =========================================================
# DOWNLOAD REPORTS
# =========================================================

if "report_data" in st.session_state:

    report_photos = st.session_state.get("report_photos", [])
    word_data = create_word_report(st.session_state["report_data"], report_photos)
    pdf_data = create_pdf_report(st.session_state["report_data"], report_photos)

    download_col1, download_col2 = st.columns(2)

    with download_col1:
        st.download_button(
            label="Download Word Report",
            data=word_data,
            file_name="quality_rejection_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with download_col2:
        st.download_button(
            label="Download PDF Report",
            data=pdf_data,
            file_name="quality_rejection_report.pdf",
            mime="application/pdf"
        )